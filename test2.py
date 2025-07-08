import streamlit as st
import os
import pandas as pd
from PyPDF2 import PdfReader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings.sentence_transformer import SentenceTransformerEmbeddings
from langchain.vectorstores import FAISS
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv
from langchain.chat_models import AzureChatOpenAI
from docx import Document

# Load environment variables
load_dotenv()

# Fetch values from .env file
AZURE_OPENAI_API_KEY = os.getenv("AZURE_OPENAI_API_KEY")
AZURE_OPENAI_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT")
api_version = os.getenv("AZURE_OPENAI_API_VERSION", "2024-02-01")

FAISS_INDEX_PATH = "faiss_index"

# Function to extract text from PDFs
def get_pdf_text(pdf_docs):
    text_chunks = []
    for pdf in pdf_docs:
        pdf_reader = PdfReader(pdf)
        for i, page in enumerate(pdf_reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                text_chunks.append((text, {"source": pdf.name, "page": i + 1}))
    return text_chunks

# Function to extract text from Excel files
def get_excel_text(excel_docs):
    chunks = []
    for excel in excel_docs:
        df = pd.read_excel(excel, sheet_name=None, dtype=str)
        for sheet_name, sheet in df.items():
            for _, row in sheet.iterrows():
                row_values = row.dropna().astype(str)
                row_text = " | ".join(row_values.tolist())
                chunks.append((row_text, {"source": excel.name, "sheet": sheet_name}))
    return chunks

# Function to extract text from TXT and DOCX files
def get_text_from_file(file):
    if file.name.endswith(".txt"):
        return [(file.read().decode("utf-8"), {"source": file.name})]
    elif file.name.endswith(".docx"):
        doc = Document(file)
        return [(para.text, {"source": file.name}) for para in doc.paragraphs if para.text.strip()]
    return []

# Function to split text into chunks
def get_text_chunks(text_chunks):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150)
    split_texts = []
    metadatas = []
    for text, metadata in text_chunks:
        splits = text_splitter.split_text(text)
        split_texts.extend(splits)
        metadatas.extend([metadata] * len(splits))
    return split_texts, metadatas

# Function to create FAISS vector store
def get_vector_store(text_chunks):
    texts, metadatas = get_text_chunks(text_chunks)
    
    # ✅ FIX: Removed `normalize_L2=True`
    embeddings = SentenceTransformerEmbeddings(model_name="all-MiniLM-L6-v2")
    
    vector_store = FAISS.from_texts(texts, embeddings, metadatas=metadatas)
    vector_store.save_local(FAISS_INDEX_PATH)
    st.success("✅ FAISS index created successfully! You can now process your questions.")

# Function to define the conversational chain
def get_conversational_chain():
    prompt_template = """
    Use the provided context to answer the question with the highest possible accuracy.
    Do not use external knowledge—only rely on the given context.
    
    Context:
    {context}
    
    Question: {question}
    Answer:
    """
    prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
    
    llm = AzureChatOpenAI(
        azure_endpoint=AZURE_OPENAI_ENDPOINT,
        api_version=api_version,
        deployment_name="gpt-35-turbo",
        api_key=AZURE_OPENAI_API_KEY
    )

    # ✅ Ensure `allow_dangerous_deserialization=True` is passed to `FAISS.load_local()`
    vector_store = FAISS.load_local(
        FAISS_INDEX_PATH, 
        SentenceTransformerEmbeddings(model_name="all-MiniLM-L6-v2"),
        allow_dangerous_deserialization=True  # Correct placement of this argument
    )
    
    retriever = vector_store.as_retriever()
    
    return RetrievalQA.from_chain_type(
        llm, 
        retriever=retriever, 
        chain_type="stuff", 
        chain_type_kwargs={"prompt": prompt}
    )


# Function to process questions from different file types
def process_questions(question_file, output_format):
    if not os.path.exists(os.path.join(FAISS_INDEX_PATH, "index.faiss")):
        st.error("No FAISS index found. Please upload and process files first.")
        return
    
    chain = get_conversational_chain()
    
    # Extract questions
    if question_file.name.endswith(".xlsx"):
        df = pd.read_excel(question_file)
        if 'Questions' not in df.columns:
            st.error("Excel file must have a 'Questions' column.")
            return
        questions = df['Questions'].tolist()
    else:
        questions = get_text_from_file(question_file)
        questions = [q[0].strip() for q in questions if q[0].strip()]
    
    answers = []
    for question in questions:
        response = chain.run(question)
        answers.append((question, response))
    
    # Save output in desired format
    if output_format == "docx":
        doc = Document()
        for q, a in answers:
            doc.add_paragraph(f"Q: {q}")
            doc.add_paragraph(f"A: {a}\n")
        output_file = "output_answers.docx"
        doc.save(output_file)
    else:
        output_file = "output_answers.txt"
        with open(output_file, "w", encoding="utf-8") as f:
            for q, a in answers:
                f.write(f"Q: {q}\nA: {a}\n\n")

    with open(output_file, "rb") as f:
        st.download_button("Download Answered Questions", f, file_name=output_file, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document" if output_format == "docx" else "text/plain")

# Streamlit App
def main():
    st.set_page_config(page_title="QnA with PDF, Excel, TXT & DOCX")
    st.header("QnA with Your Documents 💬")
    
    with st.sidebar:
        st.title("Upload Files")
        pdf_docs = st.file_uploader("Upload PDFs", type=["pdf"], accept_multiple_files=True)
        excel_docs = st.file_uploader("Upload Excel Files", type=["xls", "xlsx"], accept_multiple_files=True)
        
        if st.button("Process Documents"):
            raw_text = []
            if pdf_docs:
                raw_text.extend(get_pdf_text(pdf_docs))
            if excel_docs:
                raw_text.extend(get_excel_text(excel_docs))
            if raw_text:
                get_vector_store(raw_text)

    question_file = st.file_uploader("Upload Questions File", type=["xls", "xlsx", "txt", "docx"])
    output_format = st.selectbox("Select Output Format", ["docx", "txt"])
    if st.button("Generate Answers") and question_file:
        process_questions(question_file, output_format)

if __name__ == "__main__":
    main()
