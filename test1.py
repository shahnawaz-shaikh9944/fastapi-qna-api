import streamlit as st
import os
import pandas as pd
from PyPDF2 import PdfReader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings.sentence_transformer import SentenceTransformerEmbeddings
from langchain.vectorstores import FAISS
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv
from langchain.chat_models import AzureChatOpenAI
from docx import Document


# Load environment variables
load_dotenv()

# Fetch values from .env file
AZURE_OPENAI_API_KEY = os.getenv("AZURE_OPENAI_API_KEY")
AZURE_OPENAI_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT")
api_version = os.getenv("AZURE_OPENAI_API_VERSION", "2024-02-01")

FAISS_INDEX_PATH = "faiss_index"

# Function to extract text from PDFs
def get_pdf_text(pdf_docs):
    text = ""
    for pdf in pdf_docs:
        pdf_reader = PdfReader(pdf)
        for page in pdf_reader.pages:
            text += page.extract_text() or ""
    return text

# Function to extract text from Excel files (each row as a chunk)
def get_excel_text(excel_docs):
    chunks = []
    for excel in excel_docs:
        df = pd.read_excel(excel, sheet_name=None, dtype=str)
        for sheet in df.values():
            for _, row in sheet.iterrows():
                row_values = row.dropna().astype(str)
                row_text = " | ".join(row_values.tolist())
                chunks.append(row_text)
    return chunks

# Function to extract text from TXT and DOCX files
def get_text_from_file(file):
    if file.name.endswith(".txt"):
        return file.read().decode("utf-8")
    elif file.name.endswith(".docx"):
        doc = Document(file)
        return "\n".join([para.text for para in doc.paragraphs])
    return ""

# Function to split text into smaller chunks
def get_text_chunks(text_chunks):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=2000, chunk_overlap=200)
    return text_splitter.split_text("\n".join(text_chunks))

# Function to create FAISS vector store
def get_vector_store(text_chunks):
    embeddings = SentenceTransformerEmbeddings(model_name="all-MiniLM-L6-v2")
    vector_store = FAISS.from_texts(text_chunks, embedding=embeddings)
    vector_store.save_local(FAISS_INDEX_PATH)
    st.success("✅ FAISS index created successfully! You can now process your questions.")

# Function to define the conversational chain
def get_conversational_chain():
    prompt_template = """
    You are an AI assistant tasked with answering questions based only on the provided context.  
    Use the provided context to answer the question with the highest possible accuracy.
    Do not use external knowledge—only rely on the given context.
    
    Context:\n {context}?\n
    Question: \n{question}\n
    Answer:
    """
    prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
    
    llm = AzureChatOpenAI(
        azure_endpoint=AZURE_OPENAI_ENDPOINT,
        api_version=api_version,
        deployment_name="gpt-35-turbo",
        api_key=AZURE_OPENAI_API_KEY
    )
    
    return load_qa_chain(llm, prompt=prompt)

# Function to process questions from different file types
def process_questions(question_file, output_format):
    embeddings = SentenceTransformerEmbeddings(model_name="all-MiniLM-L6-v2")
    if not os.path.exists(os.path.join(FAISS_INDEX_PATH, "index.faiss")):
        st.error("No FAISS index found. Please upload and process files first.")
        return
    
    new_db = FAISS.load_local(FAISS_INDEX_PATH, embeddings, allow_dangerous_deserialization=True)
    chain = get_conversational_chain()

    # Extract questions based on file type
    if question_file.name.endswith(".xlsx"):
        df = pd.read_excel(question_file)
        if 'Questions' not in df.columns:
            st.error("Excel file must have a 'Questions' column.")
            return
        questions = df['Questions'].tolist()
    else:
        questions = get_text_from_file(question_file).split("\n")
        questions = [q.strip() for q in questions if q.strip()]
    
    answers = []
    for question in questions:
        docs = new_db.similarity_search(question, k=3)
        context = "\n".join([doc.page_content for doc in docs])
        answer = chain({"input_documents": docs, "question": question}, return_only_outputs=True)["output_text"]
        answers.append((question, answer))

    # Save output in desired format
    if output_format == "docx":
        doc = Document()
        for q, a in answers:
            doc.add_paragraph(f"Q: {q}")
            doc.add_paragraph(f"A: {a}\n")
        output_file = "output_answers.docx"
        doc.save(output_file)
    else:
        output_file = "output_answers.txt"
        with open(output_file, "w", encoding="utf-8") as f:
            for q, a in answers:
                   f.write(f"Q: {q}\nA: {a}\n\n")

    with open(output_file, "rb") as f:
        st.download_button("Download Answered Questions", f, file_name=output_file, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document" if output_format == "docx" else "text/plain")

# Streamlit App
def main():
    st.set_page_config(page_title="QnA with PDF, Excel, TXT & DOCX")
    st.header("QnA with Your Documents 💬")

    with st.sidebar:
        st.title("Upload Files")
        pdf_docs = st.file_uploader("Upload PDFs", type=["pdf"], accept_multiple_files=True)
        excel_docs = st.file_uploader("Upload Excel Files", type=["xls", "xlsx"], accept_multiple_files=True)
        
        if st.button("Process Documents"):
            raw_text = []
            if pdf_docs:
                raw_text.append(get_pdf_text(pdf_docs))
            if excel_docs:
                raw_text.extend(get_excel_text(excel_docs))
            if raw_text:
                text_chunks = get_text_chunks(raw_text)
                get_vector_store(text_chunks)

    question_file = st.file_uploader("Upload Questions File", type=["xls", "xlsx", "pdf", "txt", "docx"])
    output_format = st.selectbox("Select Output Format", ["docx", "txt"])
    if st.button("Generate Answers") and question_file:
        process_questions(question_file, output_format)

if __name__ == "__main__":
    main()
